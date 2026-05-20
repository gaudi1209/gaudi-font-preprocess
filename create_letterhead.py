"""
信纸模板生成器 - WPS/Word/LibreOffice 兼容
使用段落底部边框实现信纸横线效果，每行都有横线（包括空行）
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_line_border(paragraph, color="auto", size="4", space="3"):
    """为段落设置底部边框（信纸横线）"""
    pPr = paragraph._element.get_or_add_pPr()
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="%s" w:space="%s" w:color="%s"/>'
        '</w:pBdr>' % (nsdecls('w'), size, space, color)
    )
    pPr.append(pBdr)


def set_fixed_line_height(paragraph, spacing_pt=28):
    """设置固定行高（确保空行不被压缩）"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml('<w:spacing %s/>' % nsdecls('w'))
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')


def create_letterhead(output_path, num_lines=22, line_spacing=28,
                      line_color="auto", line_size="4", font_name=None,
                      font_size=None):
    """
    创建空白信纸模板

    参数:
        output_path: 输出文件路径
        num_lines: 每页行数
        line_spacing: 行距(pt), 28=标准手写, 20=打印
        line_color: 横线颜色, "auto"=黑, "99BBDD"=淡蓝
        line_size: 线条粗细(1/8磅), 4=0.5磅
        font_name: 默认字体名称
        font_size: 默认字体大小(pt)
    """
    doc = Document()

    # 页面设置 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 默认样式
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    if font_name:
        style.font.name = font_name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        style.font.size = Pt(font_size)

    # 生成信纸行
    for _ in range(num_lines):
        p = doc.add_paragraph()
        if not p.runs:
            p.add_run('')
        set_fixed_line_height(p, line_spacing)
        set_line_border(p, color=line_color, size=line_size)

    # 末尾空段落便于续写
    last = doc.add_paragraph()
    last.add_run('')
    set_fixed_line_height(last, line_spacing)
    set_line_border(last, color=line_color, size=line_size)

    doc.save(output_path)
    print(f'信纸模板已生成: {output_path} ({num_lines}行, {line_spacing}pt行距)')


def apply_letterhead_to_doc(doc_path, line_color="auto", line_size="4"):
    """
    为已有文档添加信纸横线
    将每个段落设置固定行高和底部边框
    """
    doc = Document(doc_path)

    for para in doc.paragraphs:
        # 确保有 run（空段落也需要）
        if not para.runs:
            para.add_run('')

        # 设置底部边框
        set_line_border(para, color=line_color, size=line_size)

    doc.save(doc_path)
    print(f'已为文档添加信纸横线: {doc_path}')


if __name__ == '__main__':
    output_dir = r'D:\Claudecode\高迪书法字库制作'

    # 淡蓝信纸
    create_letterhead(
        f'{output_dir}/信纸模板_淡蓝.docx',
        num_lines=22, line_spacing=28,
        line_color="99BBDD",
    )

    # 黑色信纸
    create_letterhead(
        f'{output_dir}/信纸模板_黑色.docx',
        num_lines=25, line_spacing=24,
        line_color="000000",
    )
