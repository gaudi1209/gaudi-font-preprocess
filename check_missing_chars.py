#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""比较 docx 文档与字体文件，找出缺失的字符"""

import re
from pathlib import Path
from docx import Document
from fontTools.ttLib import TTFont

# 文件路径
DOCX_PATH = r"D:\Claudecode\高迪书法字库制作\高迪字体测试20260327.docx"
FONT_PATH = r"C:\Users\chenlin\Desktop\2026工作文件\gaudi\OpenType-TT\高迪书法_王羲之V1-Regular.ttf"
OUTPUT_DIR = Path(r"D:\Claudecode\高迪书法字库制作\缺失字符")

def extract_chars_from_docx(docx_path):
    """从 docx 文件中提取所有汉字"""
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    # 提取汉字
    chars = set(re.findall(r'[\u4e00-\u9fff]', text))
    return chars

def get_font_cmap(font_path):
    """获取字体文件中包含的字符"""
    font = TTFont(font_path)
    cmap = font.getBestCmap()
    font.close()
    return set(cmap.keys())

def main():
    # 创建输出目录
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 提取 docx 中的汉字
    print(f"正在读取文档: {DOCX_PATH}")
    doc_chars = extract_chars_from_docx(DOCX_PATH)
    print(f"文档中共有 {len(doc_chars)} 个不重复汉字")

    # 获取字体中的字符
    print(f"正在读取字体: {FONT_PATH}")
    font_chars = get_font_cmap(FONT_PATH)
    print(f"字体中共有 {len(font_chars)} 个字符")

    # 找出缺失的汉字
    missing_chars = sorted([c for c in doc_chars if ord(c) not in font_chars])
    print(f"缺失 {len(missing_chars)} 个汉字")

    if missing_chars:
        # 输出缺失字符列表
        print("\n缺失字符预览（前50个）:")
        print("".join(missing_chars[:50]))

        # 每200个一组写入文件
        chunk_size = 200
        for i in range(0, len(missing_chars), chunk_size):
            chunk = missing_chars[i:i+chunk_size]
            chunk_num = i // chunk_size + 1
            output_file = OUTPUT_DIR / f"缺失字符_{chunk_num}.txt"
            with open(output_file, "w", encoding="utf-8") as f:
                f.write("".join(chunk))
            print(f"已写入: {output_file} ({len(chunk)} 个字符)")

        # 同时输出所有缺失字符到一个文件
        all_file = OUTPUT_DIR / "缺失字符_全部.txt"
        with open(all_file, "w", encoding="utf-8") as f:
            f.write("".join(missing_chars))
        print(f"已写入: {all_file}")

        # 输出 Unicode 编码列表
        unicode_file = OUTPUT_DIR / "缺失字符_Unicode.txt"
        with open(unicode_file, "w", encoding="utf-8") as f:
            for c in missing_chars:
                f.write(f"U+{ord(c):04X}  {c}\n")
        print(f"已写入: {unicode_file}")
    else:
        print("没有缺失的汉字！")

if __name__ == "__main__":
    main()
